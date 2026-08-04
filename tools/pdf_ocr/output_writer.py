"""
格式化输出器 —— 将 OCR 结果流式写入 Markdown 和 DOCX 文件。

针对 400+ 页大文档设计：
- Markdown: 流式追加（每页直接 append + flush），中断不丢失
- DOCX: 内存缓冲区积攒 50 页后批量写入，支持分卷
- 段落重排：按 Y 坐标排序 → X 坐标分行 → 相邻行合并为段落

用法:
    writer = OutputWriter("./output", "教材", split_pages=100)
    writer.write_page(1, ocr_blocks)   # 每页追加
    ...
    writer.finalize('both')            # 完成写入
    print(writer.output_paths)
"""

import logging
import os
import re
from pathlib import Path
from typing import Dict, List, Optional

from PIL import Image

logger = logging.getLogger(__name__)


class OutputWriterError(Exception):
    """输出写入异常。"""
    pass


class OutputWriter:
    """将 OCR 结果流式写入 Markdown 和 DOCX。

    流式策略：
    - MD: 每页 append + flush → 中断不丢已处理页
    - DOCX: 积攒 50 页 → document.save() → 清空缓冲区
             split_pages > 0 时每 N 页生成一个新文件
    """

    # DOCX 内部缓冲区大小（页数），超过后触发 save + 清空
    DOCX_FLUSH_INTERVAL = 50

    def __init__(
        self,
        output_dir: str,
        basename: str,
        split_pages: int = 0,
        md_enabled: bool = True,
        docx_enabled: bool = True,
    ):
        """初始化输出器。

        Args:
            output_dir: 输出目录
            basename: 输出文件基础名（不含扩展名），如 "计算机网络"
            split_pages: 每 N 页拆分为独立文件。0 表示不拆分
            md_enabled: 是否输出 Markdown
            docx_enabled: 是否输出 DOCX
        """
        self.output_dir = str(Path(output_dir).resolve())
        self.basename = basename
        self.split_pages = split_pages
        self.md_enabled = md_enabled
        self.docx_enabled = docx_enabled

        self._output_paths: List[str] = []

        # MD 文件句柄（流式写入）
        self._md_handle = None
        self._md_path: Optional[str] = None
        self._md_page_count = 0

        # DOCX 缓冲区
        self._docx_doc = None
        self._docx_buffer_pages = 0
        self._docx_current_start_page = 0
        self._docx_part_index = -1  # -1 确保首页触发 start_page 初始化
        self._docx_paths: List[str] = []

        # 确保输出目录存在
        os.makedirs(self.output_dir, exist_ok=True)

    @property
    def output_paths(self) -> List[str]:
        """返回所有已生成的输出文件路径。"""
        paths = []
        if self._md_path:
            paths.append(self._md_path)
        paths.extend(self._docx_paths)
        return paths

    def write_page(self, page_num: int, blocks: List[dict]):
        """将一页 OCR 结果写入输出。

        Args:
            page_num: 页码（1-based）
            blocks: OCR 识别结果列表，每项 {text, confidence, bbox}
        """
        # 重排文本块为段落
        paragraphs = self._reconstruct_paragraphs(blocks)

        # 流式写入 Markdown
        if self.md_enabled:
            self._write_md_page(page_num, paragraphs)

        # 缓冲写入 DOCX
        if self.docx_enabled:
            self._write_docx_page(page_num, paragraphs)

    def finalize(self, fmt: str = 'both'):
        """完成输出，关闭文件句柄，写入剩余的 DOCX 缓冲区。

        Args:
            fmt: 输出格式 — 'md', 'docx', 'both'
        """
        if fmt in ('md', 'both') and self._md_handle:
            self._md_handle.close()
            self._md_handle = None
            logger.info(f"Markdown 已保存: {self._md_path}")

        if fmt in ('docx', 'both') and self._docx_doc:
            self._flush_docx()
            logger.info(f"DOCX 文件数: {len(self._docx_paths)}")

        self._output_paths = self.output_paths

    # ------------------------------------------------------------------
    # Markdown 输出
    # ------------------------------------------------------------------

    def _write_md_page(self, page_num: int, paragraphs: List[str]):
        """流式追加一页到 Markdown 文件。"""
        # 延迟打开文件（首次写入时）
        if self._md_handle is None:
            self._md_path = os.path.join(self.output_dir, f"{self.basename}.md")
            self._md_handle = open(self._md_path, 'w', encoding='utf-8-sig')
            # 写入文件头
            self._md_handle.write(f"# {self.basename} — OCR 识别结果\n\n")
            self._md_handle.write(f"> 生成时间: {self._now()}\n\n")

        # 页码分隔
        self._md_handle.write(f"\n---\n\n### 第 {page_num} 页\n\n")

        # 写入段落
        if paragraphs:
            for para in paragraphs:
                self._md_handle.write(para + '\n\n')
        else:
            self._md_handle.write('*（本页未检测到文本）*\n\n')

        # 立即刷新到磁盘（中断不丢失）
        self._md_handle.flush()
        os.fsync(self._md_handle.fileno())

        self._md_page_count += 1

    # ------------------------------------------------------------------
    # DOCX 输出
    # ------------------------------------------------------------------

    def _write_docx_page(self, page_num: int, paragraphs: List[str]):
        """缓冲一页到 DOCX。"""
        # 检查是否需要分卷
        if self.split_pages > 0:
            current_part = (page_num - 1) // self.split_pages
            if current_part != self._docx_part_index:
                if self._docx_doc:
                    self._flush_docx()
                self._docx_part_index = current_part
                self._docx_current_start_page = current_part * self.split_pages + 1

        # 延迟创建文档
        if self._docx_doc is None:
            self._docx_doc = self._create_docx_document()

        # 添加分页符（第一页不加）
        if self._docx_buffer_pages > 0:
            self._docx_doc.add_page_break()

        # 添加页码标题
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        title_para = self._docx_doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(f"— 第 {page_num} 页 —")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(150, 150, 150)

        # 添加段落
        if paragraphs:
            for para_text in paragraphs:
                p = self._docx_doc.add_paragraph(para_text)
                # 段落间距
                pf = p.paragraph_format
                pf.space_after = Pt(6)
                pf.line_spacing = 1.5
        else:
            p = self._docx_doc.add_paragraph('（本页未检测到文本）')
            pf = p.paragraph_format
            pf.space_after = Pt(6)

        self._docx_buffer_pages += 1

        # 缓冲区达到阈值时 flush
        if self._docx_buffer_pages >= self.DOCX_FLUSH_INTERVAL:
            self._flush_docx(is_final=False)

    def _create_docx_document(self):
        """创建一个新的 DOCX 文档对象，设置默认样式。"""
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.oxml.ns import qn

        doc = Document()

        # 页面设置
        section = doc.sections[0]
        section.page_width = Cm(21.0)   # A4
        section.page_height = Cm(29.7)

        # 默认字体
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Microsoft YaHei'
        font.size = Pt(10.5)
        # 设置中文字体
        style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

        return doc

    def _flush_docx(self, is_final: bool = False):
        """将 DOCX 缓冲区写入文件。"""
        if self._docx_doc is None or self._docx_buffer_pages == 0:
            return

        # 生成文件名
        if self.split_pages > 0:
            # 分卷模式: 教材_p001-100.docx
            end_page = min(
                self._docx_current_start_page + self.split_pages - 1,
                self._docx_current_start_page + self._docx_buffer_pages - 1,
            )
            filename = (
                f"{self.basename}_p{self._docx_current_start_page:03d}-{end_page:03d}.docx"
            )
        else:
            filename = f"{self.basename}.docx"

        path = os.path.join(self.output_dir, filename)

        try:
            self._docx_doc.save(path)
            self._docx_paths.append(path)
            logger.debug(
                f"DOCX 已保存: {filename} ({self._docx_buffer_pages} 页)"
            )
        except Exception as e:
            logger.error(f"DOCX 保存失败: {e}")

        # 清空缓冲区（非最终写入时重建文档继续）
        self._docx_doc = None
        self._docx_buffer_pages = 0

        if not is_final:
            self._docx_doc = self._create_docx_document()

    # ------------------------------------------------------------------
    # 段落重排
    # ------------------------------------------------------------------

    @staticmethod
    def _reconstruct_paragraphs(blocks: List[dict]) -> List[str]:
        """将 OCR 文本块重排为段落。

        算法：
        1. 按 Y 坐标排序（从上到下）
        2. 相近 Y（< 行高阈值）的块归为同一行
        3. 行内按 X 坐标排序（从左到右）
        4. 相邻行合并为段落（Y 间距 < 段落阈值）

        Args:
            blocks: OCR 结果列表 [{text, confidence, bbox}]

        Returns:
            段落文本列表
        """
        if not blocks:
            return []

        # 1. 按 Y 坐标排序（使用 bbox 的 y1）
        sorted_blocks = sorted(blocks, key=lambda b: b['bbox'][1])

        # 2. 归并为行：相邻块 Y 坐标差 < 平均行高
        if len(sorted_blocks) >= 2:
            avg_height = sum(
                b['bbox'][3] - b['bbox'][1] for b in sorted_blocks
            ) / len(sorted_blocks)
        else:
            avg_height = sorted_blocks[0]['bbox'][3] - sorted_blocks[0]['bbox'][1]

        line_threshold = avg_height * 0.5  # Y 差小于半行高视为同一行

        lines = []  # [[block, block], [block], ...]
        current_line = [sorted_blocks[0]]

        for block in sorted_blocks[1:]:
            prev_y_center = (
                current_line[-1]['bbox'][1] + current_line[-1]['bbox'][3]
            ) / 2
            curr_y_center = (block['bbox'][1] + block['bbox'][3]) / 2

            if abs(curr_y_center - prev_y_center) < line_threshold:
                # 同一行
                current_line.append(block)
            else:
                # 新行
                lines.append(current_line)
                current_line = [block]

        lines.append(current_line)  # 最后一行

        # 每行内按 X 排序
        for line in lines:
            line.sort(key=lambda b: b['bbox'][0])

        # 3. 行 → 段落：相邻行 Y 间距 < 段落阈值
        para_threshold = avg_height * 1.8

        paragraphs = []
        current_para = [lines[0]]

        for line in lines[1:]:
            prev_y_bottom = max(b['bbox'][3] for b in current_para[-1])
            curr_y_top = min(b['bbox'][1] for b in line)

            if (curr_y_top - prev_y_bottom) < para_threshold:
                current_para.append(line)
            else:
                paragraphs.append(current_para)
                current_para = [line]

        paragraphs.append(current_para)

        # 4. 每个段落拼接为字符串
        result = []
        for para in paragraphs:
            para_lines = []
            for line in para:
                line_text = ' '.join(b['text'] for b in line)
                para_lines.append(line_text)
            result.append('\n'.join(para_lines))

        return result

    # ------------------------------------------------------------------
    # 工具方法
    # ------------------------------------------------------------------

    @staticmethod
    def _now() -> str:
        """当前时间字符串。"""
        from datetime import datetime
        return datetime.now().strftime('%Y-%m-%d %H:%M:%S')


# --- 模块自测 (python output_writer.py) ---
if __name__ == "__main__":
    import tempfile
    import shutil

    tmp_dir = tempfile.mkdtemp()

    # 模拟 OCR 结果（两段文字，分两行）
    sample_blocks = [
        # 第一段第一行
        {'text': '计算机网络', 'confidence': 0.99, 'bbox': [100, 100, 250, 130]},
        {'text': '是指', 'confidence': 0.95, 'bbox': [260, 102, 320, 128]},
        # 第一段第二行（Y 间距小 → 合并段落）
        {'text': '将多台计算机', 'confidence': 0.98, 'bbox': [100, 145, 280, 175]},
        {'text': '互联互通', 'confidence': 0.97, 'bbox': [290, 147, 380, 173]},
        # 第二段第一行（Y 间距大 → 新段落）
        {'text': 'TCP协议', 'confidence': 0.96, 'bbox': [100, 280, 200, 310]},
        {'text': '三次握手', 'confidence': 0.94, 'bbox': [210, 282, 330, 308]},
        # 空场景
    ]

    print("=== Paragraph Reorder Test ===")
    paragraphs = OutputWriter._reconstruct_paragraphs(sample_blocks)
    for i, p in enumerate(paragraphs):
        print(f"Para {i + 1}: {repr(p[:50])}...")
    assert len(paragraphs) == 2, f"Expected 2 paragraphs, got {len(paragraphs)}"
    print("[OK] Paragraph reorder correct")

    # Write test
    print("\n=== Output Write Test ===")
    writer = OutputWriter(tmp_dir, "test_output", split_pages=0)
    writer.write_page(1, sample_blocks)
    writer.write_page(2, sample_blocks)
    writer.finalize('both')

    for p in writer.output_paths:
        size = os.path.getsize(p)
        print(f"  {os.path.basename(p)}: {size} bytes")
        assert size > 0, f"File is empty: {p}"

    # Verify MD content
    with open(writer.output_paths[0], 'r', encoding='utf-8-sig') as f:
        content = f.read()
        assert '###' in content
        assert 'Page' in content or 'page' in content or '页' in content
        print("[OK] MD content verified")

    # Clean up
    shutil.rmtree(tmp_dir, ignore_errors=True)
    print("\n[OK] OutputWriter self-test passed")
