"""
资料库 API Blueprint —— 星伴(StarPal) 资料库 REST 端点。

提供双轨知识架构的后端接口:
  - 系统知识库 (user_id=NULL): 管理员上传权威教材，所有用户共享
  - 个人资料库 (user_id=具体值): 用户自主上传，仅本人可见

端点列表 (14 个):
  POST   /api/library/upload                — 上传文件
  GET    /api/library/progress/<task_id>    — SSE 进度流
  GET    /api/library/documents             — 文档列表
  DELETE /api/library/documents/<doc_id>    — 级联删除
  GET    /api/library/knowledge-coverage    — 知识点覆盖度
  GET    /api/library/knowledge-graph       — 知识图谱数据
  GET    /api/library/knowledge-node-detail — 图谱节点下钻
  POST   /api/library/analyze/<doc_id>      — 手动触发知识点映射
  GET    /api/library/learning-card         — 获取学习卡片（精简版）
  GET    /api/library/learning-card/stream  — SSE 流式生成完整版卡片
  GET    /api/library/parent-kps            — 获取标准知识点列表
  GET    /api/library/cards-status          — 批量卡片状态查询
  POST   /api/library/documents/<doc_id>/readable — 触发 AI 整理文档
  GET    /api/library/documents/<doc_id>/readable — 获取 AI 整理结果
"""

import json
import os
import threading
import time
import uuid
from pathlib import Path
from typing import Optional

from flask import Blueprint, Response, jsonify, request

# 创建 Blueprint
library_bp = Blueprint("library", __name__)

# 项目根目录
BASE_DIR = Path(__file__).parent.parent

# 上传目录（非 Web 可访问，仅 API 内部读写）
UPLOAD_FOLDER = BASE_DIR / "uploads"
SYSTEM_FOLDER = UPLOAD_FOLDER / "system"

# 允许的文件扩展名
ALLOWED_EXTENSIONS = {".pdf", ".docx"}

# 最大文件大小 (200MB)
MAX_CONTENT_LENGTH = 200 * 1024 * 1024


def _ensure_upload_dirs() -> None:
    """确保上传目录存在。"""
    os.makedirs(SYSTEM_FOLDER, exist_ok=True)


def _allowed_file(filename: str) -> bool:
    """检查文件扩展名是否在白名单。"""
    ext = os.path.splitext(filename or "")[1].lower()
    return ext in ALLOWED_EXTENSIONS


def _resolve_user_id(raw_user_id) -> Optional[int]:
    """解析请求中的 user_id 参数。
    user_id=0 或 "0" → None（系统知识库）
    其他数值 → int（个人资料库）
    """
    if raw_user_id is None:
        return None
    try:
        uid = int(raw_user_id)
        return uid if uid > 0 else None
    except (ValueError, TypeError):
        return None


def _save_uploaded_file(file, user_id: Optional[int]) -> tuple:
    """保存上传的文件到适当目录。

    Args:
        file: Flask request.files['file'] 对象。
        user_id: None=系统，其他=用户ID。

    Returns:
        (saved_path, original_filename) 元组。
    """
    _ensure_upload_dirs()

    original_name = file.filename or "unknown"
    # 生成唯一文件名以避免冲突
    file_uuid = uuid.uuid4().hex[:12]
    safe_name = f"{file_uuid}_{original_name}"

    if user_id is None:
        dest_dir = SYSTEM_FOLDER
    else:
        dest_dir = UPLOAD_FOLDER / str(user_id)
        os.makedirs(dest_dir, exist_ok=True)

    save_path = dest_dir / safe_name
    file.save(str(save_path))

    return str(save_path), original_name


def _file_size_mb(path: str) -> float:
    """获取文件大小（MB）。"""
    try:
        return os.path.getsize(path) / (1024 * 1024)
    except OSError:
        return 0.0


# ================================================================
# 端点实现
# ================================================================


@library_bp.route("/api/library/upload", methods=["POST"])
def upload_file():
    """上传资料文件（系统或用户）。

    请求: multipart/form-data
      - file: 文件（必填）
      - user_id: 0=系统知识库, 其他=个人资料库（必填）
      - doc_type: textbook/rfc/paper/note/other（可选，默认 other）
      - title: 文档标题（可选，默认取文件名）

    响应: {"success": true, "task_id": "uuid-xxx"}
    """
    if "file" not in request.files:
        return jsonify({"success": False, "message": "缺少文件"}), 400

    file = request.files["file"]
    if not file.filename:
        return jsonify({"success": False, "message": "未选择文件"}), 400

    if not _allowed_file(file.filename):
        return jsonify({
            "success": False,
            "message": f"不支持的文件格式，仅支持 PDF 和 Word (.docx)",
        }), 400

    # 解析参数
    raw_uid = request.form.get("user_id", "0")
    user_id = _resolve_user_id(raw_uid)
    doc_type = request.form.get("doc_type", "other")
    title = request.form.get("title", "") or file.filename

    # 保存文件
    try:
        saved_path, original_name = _save_uploaded_file(file, user_id)
    except Exception as e:
        return jsonify({"success": False, "message": f"文件保存失败: {str(e)}"}), 500

    # 创建任务
    task_id = uuid.uuid4().hex
    file_size = os.path.getsize(saved_path)

    # 文件类型检测（延迟导入避免循环依赖）
    from server.file_handler import detect_file_type, validate_file

    ok, err = validate_file(saved_path)
    if not ok:
        # 删除已保存的错误文件
        try:
            os.remove(saved_path)
        except OSError:
            pass
        return jsonify({"success": False, "message": err}), 400

    file_type = detect_file_type(saved_path) or "text_pdf"

    # 写入 library_tasks 记录
    from database.db_connector import get_connection

    conn = get_connection()
    if not conn:
        try:
            os.remove(saved_path)
        except OSError:
            pass
        return jsonify({"success": False, "message": "数据库连接失败"}), 500

    try:
        cursor = conn.cursor()
        cursor.execute(
            """INSERT INTO library_tasks
               (task_id, user_id, file_name, file_path, file_type, file_size_bytes, status)
               VALUES (%s, %s, %s, %s, %s, %s, 'pending')""",
            (task_id, user_id, original_name, saved_path, file_type, file_size),
        )
        conn.commit()
    except Exception as e:
        conn.rollback()
        try:
            os.remove(saved_path)
        except OSError:
            pass
        return jsonify({"success": False, "message": f"任务创建失败: {str(e)}"}), 500
    finally:
        cursor.close()
        conn.close()

    # 启动后台处理线程
    thread = threading.Thread(
        target=_process_in_background,
        args=(task_id, user_id, saved_path, original_name, file_type, doc_type, title),
        daemon=True,
    )
    thread.start()

    return jsonify({
        "success": True,
        "task_id": task_id,
        "message": "文件已上传，正在后台处理",
    })


@library_bp.route("/api/library/progress/<task_id>", methods=["GET"])
def get_progress(task_id: str):
    """SSE 流式推送任务处理进度。

    响应: text/event-stream
      data: {"progress_pct": N, "detail": {...}}
      data: [DONE]
    """
    def generate():
        from database.db_connector import get_connection

        # 立即发送初始状态
        try:
            conn = get_connection()
            if conn:
                cursor = conn.cursor(dictionary=True)
                cursor.execute(
                    "SELECT status, progress_pct, progress_detail FROM library_tasks WHERE task_id=%s",
                    (task_id,),
                )
                row = cursor.fetchone()
                if row:
                    detail = {}
                    detail_raw = row.get("progress_detail")
                    if detail_raw:
                        try:
                            detail = json.loads(detail_raw) if isinstance(detail_raw, str) else detail_raw
                        except Exception:
                            detail = {}
                    yield f"data: {json.dumps({'progress_pct': row['progress_pct'] or 0, 'status': row['status'], 'detail': detail})}\n\n"
                cursor.close()
                conn.close()
        except Exception:
            pass

        last_pct = -1.0
        finished = False
        max_polls = 3600  # 最多轮询 30 分钟 (3600 × 0.5s)，覆盖 400 页 OCR

        for _ in range(max_polls):
            conn = get_connection()
            if not conn:
                time.sleep(0.5)
                continue

            try:
                cursor = conn.cursor(dictionary=True)
                cursor.execute(
                    """SELECT status, progress_pct, progress_detail, doc_id
                       FROM library_tasks WHERE task_id = %s""",
                    (task_id,),
                )
                row = cursor.fetchone()
            finally:
                cursor.close()
                conn.close()

            if row is None:
                yield f"data: {json.dumps({'error': '任务不存在'})}\n\n"
                yield "data: [DONE]\n\n"
                return

            pct = row["progress_pct"] or 0.0
            status = row["status"]
            detail_raw = row.get("progress_detail")

            # 解析 detail JSON
            detail = {}
            if detail_raw:
                try:
                    detail = json.loads(detail_raw) if isinstance(detail_raw, str) else detail_raw
                except (json.JSONDecodeError, TypeError):
                    detail = {"raw": str(detail_raw)}

            # 仅在进度变化时推送
            if abs(pct - last_pct) > 0.5 or status in ("completed", "failed"):
                payload = {
                    "progress_pct": pct,
                    "status": status,
                    "detail": detail,
                }
                if row.get("doc_id"):
                    payload["doc_id"] = row["doc_id"]
                yield f"data: {json.dumps(payload)}\n\n"
                last_pct = pct

            if status in ("completed", "failed"):
                finished = True
                break

            time.sleep(0.5)

        if finished:
            yield "data: [DONE]\n\n"
        else:
            yield f"data: {json.dumps({'error': '任务超时'})}\n\n"
            yield "data: [DONE]\n\n"

    return Response(
        generate(),
        mimetype="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


@library_bp.route("/api/library/documents", methods=["GET"])
def list_documents():
    """获取文档列表（系统或用户）。

    Query 参数:
      user_id: 0=系统知识库, 其他=个人资料库（必填）
      page: 页码（默认 1）
      page_size: 每页条数（默认 20）
    """
    raw_uid = request.args.get("user_id", "0")
    user_id = _resolve_user_id(raw_uid)

    page = max(1, int(request.args.get("page", 1) or 1))
    page_size = min(100, max(1, int(request.args.get("page_size", 20) or 20)))
    offset = (page - 1) * page_size

    from database.db_connector import get_connection

    conn = get_connection()
    if not conn:
        return jsonify({"success": False, "message": "数据库连接失败"}), 500

    try:
        cursor = conn.cursor(dictionary=True)

        # 构建过滤条件
        if user_id is None:
            where = "WHERE kd.user_id IS NULL"
            count_where = "WHERE user_id IS NULL"
        else:
            where = "WHERE kd.user_id = %s"
            count_where = "WHERE user_id = %s"

        # 总数
        cursor.execute(
            f"SELECT COUNT(*) AS cnt FROM knowledge_documents {count_where}",
            (user_id,) if user_id is not None else (),
        )
        total = (cursor.fetchone() or {}).get("cnt", 0)

        # 文档列表（联表获取 chunk_count 和最新任务状态）
        sql = f"""
            SELECT
                kd.doc_id,
                kd.title,
                kd.doc_type,
                kd.status,
                kd.created_at,
                COUNT(kc.chunk_id) AS chunk_count,
                (SELECT lt.status FROM library_tasks lt
                 WHERE lt.doc_id = kd.doc_id ORDER BY lt.updated_at DESC LIMIT 1
                ) AS task_status
            FROM knowledge_documents kd
            LEFT JOIN knowledge_chunks kc ON kd.doc_id = kc.doc_id
            {where}
            GROUP BY kd.doc_id
            ORDER BY kd.created_at DESC
            LIMIT %s OFFSET %s
        """
        params = (user_id,) if user_id is not None else ()
        params = params + (page_size, offset)
        cursor.execute(sql, params)
        rows = cursor.fetchall() or []

        documents = []
        for r in rows:
            documents.append({
                "doc_id": r["doc_id"],
                "title": r["title"],
                "doc_type": r["doc_type"],
                "status": r["status"],
                "task_status": r.get("task_status", ""),
                "chunk_count": r["chunk_count"] or 0,
                "created_at": str(r["created_at"]) if r.get("created_at") else "",
            })

        return jsonify({
            "success": True,
            "documents": documents,
            "total": total,
            "page": page,
            "page_size": page_size,
        })

    except Exception as e:
        return jsonify({"success": False, "message": f"查询失败: {str(e)}"}), 500
    finally:
        cursor.close()
        conn.close()


@library_bp.route("/api/library/documents/<int:doc_id>", methods=["DELETE"])
def delete_document(doc_id: int):
    """级联删除文档及其关联数据。

    删除范围: 物理文件 + MySQL 知识块(外键 CASCADE) + ChromaDB 向量 + library_tasks 记录

    Query 参数:
      user_id: 用于权限校验（可选，不传则不校验）
    """
    raw_uid = request.args.get("user_id")
    user_id = _resolve_user_id(raw_uid) if raw_uid else None

    from database.db_connector import get_connection

    conn = get_connection()
    if not conn:
        return jsonify({"success": False, "message": "数据库连接失败"}), 500

    try:
        cursor = conn.cursor(dictionary=True)

        # 获取文档信息
        cursor.execute(
            "SELECT doc_id, title, user_id FROM knowledge_documents WHERE doc_id = %s",
            (doc_id,),
        )
        doc = cursor.fetchone()
        if not doc:
            return jsonify({"success": False, "message": "文档不存在"}), 404

        # 权限校验：仅文档所有者或系统文档（user_id=NULL）可删除
        # 系统文档由管理员删除（不校验 user_id 参数时）
        if doc["user_id"] is not None and user_id is not None:
            if doc["user_id"] != user_id:
                return jsonify({"success": False, "message": "无权操作此文档"}), 403

        # 获取关联的 chunk_id 列表（用于清理 ChromaDB）
        cursor.execute(
            "SELECT chunk_id FROM knowledge_chunks WHERE doc_id = %s",
            (doc_id,),
        )
        chunk_ids = [r["chunk_id"] for r in (cursor.fetchall() or [])]

        # 清理 ChromaDB
        if chunk_ids:
            try:
                from AI_operate.rag_service import RAGService
                rag = RAGService()
                rag.delete_chunks(chunk_ids)
            except Exception as e:
                print(f"[library_api] ChromaDB 清理失败: {e}")

        # 清理 library_tasks 记录
        cursor.execute(
            "DELETE FROM library_tasks WHERE doc_id = %s", (doc_id,)
        )

        # 清理物理文件（从 library_tasks 获取文件路径）
        cursor.execute(
            "SELECT file_path FROM library_tasks WHERE doc_id = %s LIMIT 1",
            (doc_id,),
        )
        task = cursor.fetchone()
        if task and task.get("file_path"):
            try:
                os.remove(task["file_path"])
            except OSError:
                pass

        # 删除 MySQL 记录（knowledge_chunks 由 ON DELETE CASCADE 自动删除）
        cursor.execute(
            "DELETE FROM knowledge_documents WHERE doc_id = %s", (doc_id,)
        )
        conn.commit()

        return jsonify({
            "success": True,
            "message": f"文档 '{doc['title']}' 及其 {len(chunk_ids)} 个知识块已删除",
        })

    except Exception as e:
        conn.rollback()
        return jsonify({"success": False, "message": f"删除失败: {str(e)}"}), 500
    finally:
        cursor.close()
        conn.close()


@library_bp.route("/api/library/knowledge-coverage", methods=["GET"])
def get_coverage():
    """获取知识点覆盖度。

    Query 参数:
      user_id: 0=系统知识库, 其他=个人资料库（必填）
    """
    raw_uid = request.args.get("user_id", "0")
    user_id = _resolve_user_id(raw_uid)

    from server.knowledge_mapper import KnowledgeMapper

    mapper = KnowledgeMapper()
    result = mapper.get_user_coverage(user_id=user_id)

    if "error" in result:
        return jsonify({"success": False, "message": result["error"]}), 500

    return jsonify({"success": True, "coverage": result})


@library_bp.route("/api/library/knowledge-graph", methods=["GET"])
def get_graph():
    """获取知识图谱数据。

    Query 参数:
      user_id: 0=系统知识库, 其他=个人资料库（必填）
    """
    raw_uid = request.args.get("user_id", "0")
    user_id = _resolve_user_id(raw_uid)

    from server.knowledge_mapper import KnowledgeMapper

    mapper = KnowledgeMapper()
    result = mapper.get_knowledge_graph(user_id=user_id)

    if "error" in result:
        return jsonify({"success": False, "message": result["error"]}), 500

    return jsonify({"success": True, "graph": result})


@library_bp.route("/api/library/knowledge-node-detail", methods=["GET"])
def get_node_detail():
    """获取知识图谱节点的下钻详情。

    Query 参数:
      sub_topic_id: 子知识点 ID（必填）
      user_id: 0=系统知识库, 其他=个人资料库（必填）
    """
    sub_topic_id = request.args.get("sub_topic_id")
    if not sub_topic_id:
        return jsonify({"success": False, "message": "缺少 sub_topic_id 参数"}), 400

    try:
        sub_topic_id = int(sub_topic_id)
    except (ValueError, TypeError):
        return jsonify({"success": False, "message": "sub_topic_id 必须为整数"}), 400

    raw_uid = request.args.get("user_id", "0")
    user_id = _resolve_user_id(raw_uid)

    from server.knowledge_mapper import KnowledgeMapper

    mapper = KnowledgeMapper()
    result = mapper.get_node_detail(sub_topic_id, user_id=user_id)

    if "error" in result:
        return jsonify({"success": False, "message": result["error"]}), 500

    return jsonify({"success": True, "node_detail": result})


@library_bp.route("/api/library/analyze/<int:doc_id>", methods=["POST"])
def analyze_document(doc_id: int):
    """手动触发单个文档的知识点映射（处理管线会自动调用，此端点作为手动重试）。

    路径参数:
      doc_id: 文档 ID
    """
    from server.knowledge_mapper import KnowledgeMapper

    mapper = KnowledgeMapper()
    result = mapper.analyze_document(doc_id)

    if "error" in result:
        return jsonify({"success": False, "message": result["error"]}), 500

    return jsonify({"success": True, "analysis": result})


@library_bp.route("/api/library/documents/<int:doc_id>/export", methods=["GET"])
def export_document(doc_id: int):
    """导出文档全部文本。

    支持三种格式（通过 ?format= 参数指定）：
      - md (默认): Markdown 文本
      - docx: Word 文档
      - pdf: 文字版 PDF

    将所有 chunks 按 chunk_index 排序拼接后导出。
    """
    from database.db_connector import get_connection

    export_format = request.args.get("format", "md").lower()
    if export_format not in ("md", "docx", "pdf"):
        return jsonify({"success": False, "message": "不支持的导出格式，可选：md, docx, pdf"}), 400

    conn = get_connection()
    if not conn:
        return jsonify({"success": False, "message": "数据库连接失败"}), 500

    try:
        cursor = conn.cursor(dictionary=True)
        cursor.execute(
            "SELECT title FROM knowledge_documents WHERE doc_id = %s", (doc_id,)
        )
        doc = cursor.fetchone()
        if not doc:
            return jsonify({"success": False, "message": "文档不存在"}), 404

        title = doc["title"]
        cursor.execute(
            "SELECT chunk_index, content FROM knowledge_chunks "
            "WHERE doc_id = %s ORDER BY chunk_index",
            (doc_id,),
        )
        chunks = cursor.fetchall() or []

        from urllib.parse import quote
        safe_title = quote(title, safe='')

        if export_format == "md":
            return _export_markdown(title, chunks, safe_title)
        elif export_format == "docx":
            return _export_docx(title, chunks, safe_title)
        elif export_format == "pdf":
            return _export_pdf(title, chunks, safe_title)

    except Exception as e:
        return jsonify({"success": False, "message": f"导出失败: {str(e)}"}), 500
    finally:
        cursor.close()
        conn.close()


def _export_markdown(title: str, chunks: list, safe_title: str) -> Response:
    """导出为 Markdown 格式。"""
    lines = [f"# {title}\n"]
    for ch in chunks:
        lines.append(ch["content"])
        lines.append("")

    md_text = "\n\n".join(lines)
    resp = Response(
        md_text.encode("utf-8"),
        content_type="text/markdown; charset=utf-8",
    )
    resp.headers["Content-Disposition"] = f"attachment; filename*=UTF-8''{safe_title}.md"
    resp.headers["Cache-Control"] = "no-cache"
    return resp


def _export_docx(title: str, chunks: list, safe_title: str) -> Response:
    """导出为 Word (.docx) 格式。"""
    import io as io_module
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(11)

    # 标题
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 正文内容
    for ch in chunks:
        content = ch["content"].strip()
        if not content:
            continue
        # 检测是否看起来像标题（短行、以第X章/第X节开头等）
        if len(content) < 60 and (
            content.startswith("第") or content.startswith("#") or
            content.startswith("目") or content.startswith("附")
        ):
            doc.add_heading(content, level=2)
        else:
            para = doc.add_paragraph(content)
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.line_spacing = 1.5

    # 写入 BytesIO
    buffer = io_module.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    resp = Response(
        buffer.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    resp.headers["Content-Disposition"] = f"attachment; filename*=UTF-8''{safe_title}.docx"
    resp.headers["Cache-Control"] = "no-cache"
    return resp


def _export_pdf(title: str, chunks: list, safe_title: str) -> Response:
    """导出为文字版 PDF 格式（支持中文）。"""
    import io as io_module
    from fpdf import FPDF

    # 使用系统中文字体
    _CN_FONT_PATH = "C:/Windows/Fonts/simhei.ttf"

    class PDF(FPDF):
        def header(self):
            if not self.page_no():
                return
            self.set_font("CnFont", "", 9)
            self.set_text_color(128, 128, 128)
            self.cell(0, 8, title, align="C")
            self.ln(12)

        def footer(self):
            self.set_y(-15)
            self.set_font("CnFont", "", 8)
            self.set_text_color(128, 128, 128)
            self.cell(0, 10, f"\u7b2c {self.page_no()}/{{nb}} \u9875", align="C")

    pdf = PDF()
    pdf.add_font("CnFont", "", _CN_FONT_PATH, uni=True)
    pdf.alias_nb_pages()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    # 标题页
    pdf.set_font("CnFont", "", 22)
    pdf.ln(30)
    pdf.multi_cell(0, 14, title, align="C")
    pdf.ln(10)
    pdf.set_font("CnFont", "", 10)
    pdf.set_text_color(128, 128, 128)
    pdf.cell(0, 8, f"\u5171 {len(chunks)} \u4e2a\u77e5\u8bc6\u5757", align="C")
    pdf.ln(20)

    # 正文内容
    pdf.set_text_color(0, 0, 0)
    for ch in chunks:
        content = ch["content"].strip()
        if not content:
            continue
        # 短行当小标题
        if len(content) < 60 and (
            content.startswith("\u7b2c") or content.startswith("#") or
            content.startswith("\u76ee") or content.startswith("\u9644")
        ):
            pdf.set_font("CnFont", "", 14)
            pdf.ln(4)
            pdf.multi_cell(0, 9, content)
            pdf.ln(2)
            pdf.set_font("CnFont", "", 11)
        else:
            pdf.set_font("CnFont", "", 11)
            pdf.multi_cell(0, 7, content)
            pdf.ln(2)

    # 输出
    pdf_output = pdf.output()
    resp = Response(
        pdf_output,
        content_type="application/pdf",
    )
    resp.headers["Content-Disposition"] = f"attachment; filename*=UTF-8''{safe_title}.pdf"
    resp.headers["Cache-Control"] = "no-cache"
    return resp


# ================================================================
# 学习卡片端点（v0.3.x 新增）
# ================================================================


@library_bp.route("/api/library/learning-card", methods=["GET"])
def get_learning_card():
    """获取知识点学习卡片（精简版，同步生成）。

    首次访问时同步生成精简版（~5-10s），后续从缓存读取。

    Query 参数:
      sub_topic_id: 子知识点 ID（必填）
      user_id: 0=系统知识库, 其他=个人资料库（必填）
    """
    sub_topic_id = request.args.get("sub_topic_id")
    if not sub_topic_id:
        return jsonify({"success": False, "message": "缺少 sub_topic_id 参数"}), 400

    try:
        sub_topic_id = int(sub_topic_id)
    except (ValueError, TypeError):
        return jsonify({"success": False, "message": "sub_topic_id 必须为整数"}), 400

    raw_uid = request.args.get("user_id", "0")
    user_id = _resolve_user_id(raw_uid)

    from server.learning_card_service import LearningCardService

    svc = LearningCardService()
    card = svc.get_card(sub_topic_id, user_id)

    if "error" in card:
        return jsonify({"success": False, "message": card["error"]}), 404

    return jsonify({"success": True, "card": card})


@library_bp.route("/api/library/learning-card/stream", methods=["GET"])
def stream_learning_card():
    """SSE 流式生成完整版学习卡片。

    Query 参数:
      sub_topic_id: 子知识点 ID（必填）
      user_id: 0=系统知识库, 其他=个人资料库（必填）
    """
    sub_topic_id = request.args.get("sub_topic_id")
    if not sub_topic_id:
        def _err():
            yield f"data: {json.dumps({'error': '缺少 sub_topic_id 参数'})}\n\n"
            yield "data: [DONE]\n\n"
        return Response(_err(), mimetype="text/event-stream",
                        headers={"Cache-Control": "no-cache", "Connection": "keep-alive"})

    try:
        sub_topic_id = int(sub_topic_id)
    except (ValueError, TypeError):
        def _err2():
            yield f"data: {json.dumps({'error': 'sub_topic_id 必须为整数'})}\n\n"
            yield "data: [DONE]\n\n"
        return Response(_err2(), mimetype="text/event-stream",
                        headers={"Cache-Control": "no-cache", "Connection": "keep-alive"})

    raw_uid = request.args.get("user_id", "0")
    user_id = _resolve_user_id(raw_uid)

    from server.learning_card_service import LearningCardService

    svc = LearningCardService()

    return Response(
        svc.generate_full_card_stream(sub_topic_id, user_id),
        mimetype="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


@library_bp.route("/api/library/parent-kps", methods=["GET"])
def get_parent_kps():
    """获取知识库中所有标准知识点名称（parent_kp）列表。

    从 knowledge_sub_topics 表中提取去重的 parent_kp 值。
    用于前端知识点筛选、题库对齐等场景。
    """
    from database.db_connector import get_connection

    conn = get_connection()
    if not conn:
        return jsonify({"success": False, "message": "数据库连接失败"}), 500

    try:
        cursor = conn.cursor(dictionary=True)
        cursor.execute(
            "SELECT DISTINCT parent_kp, COUNT(*) AS sub_count "
            "FROM knowledge_sub_topics "
            "GROUP BY parent_kp ORDER BY parent_kp"
        )
        rows = cursor.fetchall() or []
        parent_kps = [
            {"name": r["parent_kp"], "sub_topic_count": r["sub_count"]}
            for r in rows
        ]
        return jsonify({"success": True, "parent_kps": parent_kps, "total": len(parent_kps)})
    except Exception as e:
        return jsonify({"success": False, "message": f"查询失败: {str(e)}"}), 500
    finally:
        cursor.close()
        conn.close()


@library_bp.route("/api/library/cards-status", methods=["GET"])
def get_cards_status():
    """批量获取所有知识点的卡片生成状态。

    用于前端展示哪些知识点已有学习卡片、哪些需要生成。

    Query 参数:
      user_id: 0=系统知识库, 其他=个人资料库（必填）
    """
    raw_uid = request.args.get("user_id", "0")
    user_id = _resolve_user_id(raw_uid)

    from database.db_connector import get_connection

    conn = get_connection()
    if not conn:
        return jsonify({"success": False, "message": "数据库连接失败"}), 500

    try:
        cursor = conn.cursor(dictionary=True)

        # 获取所有已缓存的卡片状态
        if user_id is not None:
            cursor.execute(
                """SELECT sub_topic_id, slim_content IS NOT NULL AS has_slim,
                          full_content IS NOT NULL AS has_full,
                          is_regenerating, generated_at
                   FROM knowledge_learning_cards
                   WHERE user_id = %s""",
                (user_id,),
            )
        else:
            cursor.execute(
                """SELECT sub_topic_id, slim_content IS NOT NULL AS has_slim,
                          full_content IS NOT NULL AS has_full,
                          is_regenerating, generated_at
                   FROM knowledge_learning_cards
                   WHERE user_id IS NULL""",
            )

        card_rows = {
            r["sub_topic_id"]: {
                "has_slim": bool(r["has_slim"]),
                "has_full": bool(r["has_full"]),
                "is_regenerating": bool(r.get("is_regenerating", 0)),
                "generated_at": str(r["generated_at"]) if r.get("generated_at") else "",
            }
            for r in (cursor.fetchall() or [])
        }

        # 获取全部子知识点
        cursor.execute(
            "SELECT sub_topic_id, sub_topic_name, parent_kp FROM knowledge_sub_topics ORDER BY sub_topic_id"
        )
        all_topics = cursor.fetchall() or []

        statuses = []
        for t in all_topics:
            tid = t["sub_topic_id"]
            cached = card_rows.get(tid, {})
            statuses.append({
                "sub_topic_id": tid,
                "sub_topic_name": t["sub_topic_name"],
                "parent_kp": t["parent_kp"],
                "has_card": cached.get("has_slim", False),
                "has_full": cached.get("has_full", False),
                "is_regenerating": cached.get("is_regenerating", False),
                "generated_at": cached.get("generated_at", ""),
            })

        # 统计
        covered = sum(1 for s in statuses if s["has_card"])
        full_count = sum(1 for s in statuses if s["has_full"])

        return jsonify({
            "success": True,
            "statuses": statuses,
            "total_topics": len(statuses),
            "covered_topics": covered,
            "full_card_topics": full_count,
        })

    except Exception as e:
        return jsonify({"success": False, "message": f"查询失败: {str(e)}"}), 500
    finally:
        cursor.close()
        conn.close()


# ================================================================
# 文档 AI 可读内容端点（v0.3.x 新增）
# ================================================================


@library_bp.route("/api/library/documents/<int:doc_id>/readable", methods=["POST"])
def generate_readable(doc_id: int):
    """触发 AI 整理文档为可读 Markdown。

    AI 将 OCR 分块重新组织为结构化 Markdown，包括:
      - 章节目录检测
      - 表格/图片占位描述
      - 段落重组

    结果缓存到 knowledge_documents.readable_content 字段。

    请求体（可选）:
      {"force": true}  — 强制重新生成（忽略已有缓存）
    """
    from database.db_connector import get_connection

    conn = get_connection()
    if not conn:
        return jsonify({"success": False, "message": "数据库连接失败"}), 500

    try:
        cursor = conn.cursor(dictionary=True)

        # 检查文档存在
        cursor.execute(
            "SELECT doc_id, title, readable_content, readable_generated_at FROM knowledge_documents WHERE doc_id = %s",
            (doc_id,),
        )
        doc = cursor.fetchone()
        if not doc:
            return jsonify({"success": False, "message": "文档不存在"}), 404

        force = request.json.get("force", False) if request.json else False

        # 已有缓存且不强制重新生成
        if doc.get("readable_content") and not force:
            return jsonify({
                "success": True,
                "doc_id": doc_id,
                "readable_content": doc["readable_content"],
                "generated_at": str(doc.get("readable_generated_at", "")),
                "from_cache": True,
                "message": "已有 AI 整理缓存",
            })

        # 获取所有 chunk 内容
        cursor.execute(
            "SELECT chunk_index, content FROM knowledge_chunks WHERE doc_id = %s ORDER BY chunk_index",
            (doc_id,),
        )
        chunks = cursor.fetchall() or []
        if not chunks:
            return jsonify({"success": False, "message": "文档无内容块"}), 400

        # 拼接原文（限制长度防超 token）
        raw_text_parts = []
        total_chars = 0
        max_input = 30000
        for ch in chunks:
            content = ch["content"]
            if total_chars + len(content) > max_input:
                raw_text_parts.append(content[:max_input - total_chars] + "\n\n(内容过长，已截断)")
                break
            raw_text_parts.append(content)
            total_chars += len(content)

        raw_full = "\n\n".join(raw_text_parts)

        # 调用 AI 整理
        from AI_operate.deepseek_chat import deepseek_chat as _dc

        prompt = f"""你是一位专业的文档整理专家。请将以下OCR识别/分块后的计算机网络资料重新整理为结构清晰、易于阅读的Markdown文档。

## 原始内容
{raw_full}

## 整理要求
1. 检测并标注章节目录（如检测到"第X章"、"第X节"等模式，请使用Markdown标题）
2. 合并语义相关的相邻段落，确保逻辑连贯
3. 对于缺失的图片/表格，用 **[此处为图片/表格：描述]** 占位
4. 保留所有重要概念和关键术语
5. 修复明显的OCR错误（如错别字、断句错误）
6. 在文档开头生成一个简短的目录（如果有多章）
7. 输出纯Markdown格式，不要包含代码块标记```

请直接输出整理后的Markdown内容："""

        readable = _dc.chat_with_deepseek(prompt)

        # 保存到数据库
        cursor.execute(
            """UPDATE knowledge_documents
               SET readable_content = %s, readable_generated_at = NOW()
               WHERE doc_id = %s""",
            (readable, doc_id),
        )
        conn.commit()

        return jsonify({
            "success": True,
            "doc_id": doc_id,
            "readable_content": readable,
            "from_cache": False,
            "message": "AI 整理完成",
        })

    except Exception as e:
        return jsonify({"success": False, "message": f"AI 整理失败: {str(e)}"}), 500
    finally:
        cursor.close()
        conn.close()


@library_bp.route("/api/library/documents/<int:doc_id>/readable", methods=["GET"])
def get_readable(doc_id: int):
    """获取文档的 AI 整理可读内容。

    如果尚未生成，返回 404 并提示先调用 POST 生成。
    """
    from database.db_connector import get_connection

    conn = get_connection()
    if not conn:
        return jsonify({"success": False, "message": "数据库连接失败"}), 500

    try:
        cursor = conn.cursor(dictionary=True)
        cursor.execute(
            "SELECT doc_id, title, readable_content, readable_generated_at "
            "FROM knowledge_documents WHERE doc_id = %s",
            (doc_id,),
        )
        doc = cursor.fetchone()
        if not doc:
            return jsonify({"success": False, "message": "文档不存在"}), 404

        if not doc.get("readable_content"):
            return jsonify({
                "success": False,
                "message": "尚未生成可读内容，请先调用 POST 触发 AI 整理",
                "ready": False,
            }), 404

        return jsonify({
            "success": True,
            "doc_id": doc_id,
            "title": doc["title"],
            "readable_content": doc["readable_content"],
            "generated_at": str(doc.get("readable_generated_at", "")),
            "ready": True,
        })

    except Exception as e:
        return jsonify({"success": False, "message": f"查询失败: {str(e)}"}), 500
    finally:
        cursor.close()
        conn.close()


# ================================================================
# 内部函数
# ================================================================


def _process_in_background(
    task_id: str,
    user_id: Optional[int],
    file_path: str,
    file_name: str,
    file_type: str,
    doc_type: str,
    title: str,
) -> None:
    """后台线程入口：执行完整的资料处理管线。

    Args 同 LibraryPipeline.process_upload。
    """
    from server.library_pipeline import LibraryPipeline

    pipeline = LibraryPipeline()
    result = pipeline.process_upload(
        task_id=task_id,
        user_id=user_id,
        file_path=file_path,
        file_name=file_name,
        file_type=file_type,
        doc_type=doc_type,
        title=title,
    )
    print(f"[library_api] 后台任务 {task_id} 完成: {result.get('message')}")


# ================================================================
# Blueprint 注册配置
# ================================================================

def init_library(app):
    """在 Flask 应用上注册资料库 Blueprint 并配置上传参数。

    Args:
        app: Flask 应用实例。
    """
    # 上传文件大小限制
    app.config["MAX_CONTENT_LENGTH"] = MAX_CONTENT_LENGTH
    # 上传目录
    app.config["UPLOAD_FOLDER"] = str(UPLOAD_FOLDER)
    _ensure_upload_dirs()

    app.register_blueprint(library_bp)
    print("[library_api] 资料库 Blueprint 已注册 (14 个端点)")
